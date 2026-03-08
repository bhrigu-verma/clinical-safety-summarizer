"""
src/data_processing/pdf_extractor.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
REWRITTEN (v2) — Multi-Modal Clinical Table Extractor

Supports:
  • PDF  → pdfplumber (text-based) + PyMuPDF fallback (scanned PDFs)
  • DOCX → python-docx table parser
  • JPEG/PNG → delegates to image_extractor.py

Key improvements over v1:
  - Handles scanned/image-based PDFs via PyMuPDF page-to-image rendering
  - Better table-summary pair detection (uses more regex patterns)
  - Returns rich metadata per pair (table type, arm count, etc.)
  - DOCX support for .docx input files
  - Clean linearization that the KNN + slot-fill pipeline can consume
"""

import re
import json
import logging
import tempfile
from pathlib import Path
from typing import List, Dict, Optional, Tuple

logger = logging.getLogger(__name__)


class ClinicalTablePair:
    """One extracted (table, writeup) pair."""
    def __init__(
        self,
        table_text: str,      # linearized format
        writeup: str,         # reference narrative
        source: str = "",
        table_type: str = "unknown",
        page: int = 0,
        pair_id: str = ""
    ):
        self.table_text = table_text
        self.writeup = writeup
        self.source = source
        self.table_type = table_type
        self.page = page
        self.pair_id = pair_id

    def to_dict(self) -> Dict:
        return {
            "table_text": self.table_text,
            "writeup": self.writeup,
            "source": self.source,
            "table_type": self.table_type,
            "page": self.page,
            "pair_id": self.pair_id,
        }

    # Legacy keys for backward-compat with the existing ML pipeline
    @property
    def input(self) -> str:
        return self.table_text

    @property
    def output(self) -> str:
        return self.writeup


class ClinicalPDFExtractor:
    """
    Extract clinical safety table-writeup pairs from PDF documents.

    Handles both text-based PDFs (pdfplumber) and scanned PDFs (PyMuPDF
    renders each page to an image then delegates to ClinicalImageExtractor).

    Usage:
        extractor = ClinicalPDFExtractor("path/to/study.pdf")
        pairs = extractor.extract_all()   # → List[ClinicalTablePair]
        extractor.save("data/processed/")
    """

    # Patterns that signal the start of a clinical narrative following a table
    WRITEUP_STARTERS = [
        r'Table\s+\d+[\.:]\s*(?:presents|shows|displays|summarizes|provides|lists)',
        r'(?:Overall|The following|The table)',
        r'(?:Treatment-emergent|Treatment emergent|TEAE)',
        r'(?:Adverse events?|AEs?|SAEs?)\s+(?:were|are|occurred)',
        r'(?:All|Most|Nearly all)\s+subjects',
    ]
    _WRITEUP_RE = re.compile(
        '(' + '|'.join(WRITEUP_STARTERS) + ')',
        re.IGNORECASE
    )

    def __init__(self, source_path: str, use_gpu_ocr: bool = False):
        self.source_path = Path(source_path)
        self.use_gpu_ocr = use_gpu_ocr
        self.pairs: List[ClinicalTablePair] = []

    # ─────────────────────────────────────────────────────────────────────────
    # Public API
    # ─────────────────────────────────────────────────────────────────────────

    def extract_all(self) -> List[ClinicalTablePair]:
        """Route to the correct extractor based on file extension."""
        suffix = self.source_path.suffix.lower()

        if suffix == ".pdf":
            self.pairs = self._extract_pdf()
        elif suffix in (".docx", ".doc"):
            self.pairs = self._extract_docx()
        elif suffix in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"):
            self.pairs = self._extract_image()
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            self.pairs = []

        logger.info(f"Extracted {len(self.pairs)} table-writeup pairs from {self.source_path.name}")
        return self.pairs

    def save(self, output_dir: str, filename: str = "raw_pairs.json"):
        """Save all pairs to JSON."""
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        dest = out / filename
        with open(dest, "w") as f:
            json.dump([p.to_dict() for p in self.pairs], f, indent=2)
        logger.info(f"Saved {len(self.pairs)} pairs → {dest}")

    # ─────────────────────────────────────────────────────────────────────────
    # PDF Extraction
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_pdf(self) -> List[ClinicalTablePair]:
        pairs = []

        try:
            import pdfplumber
            with pdfplumber.open(str(self.source_path)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_pairs = self._process_pdf_page(page, page_num)
                    pairs.extend(page_pairs)

            if not pairs:
                logger.info("pdfplumber found no pairs; attempting scanned-PDF fallback")
                pairs = self._extract_pdf_scanned()

        except Exception as e:
            logger.warning(f"pdfplumber failed ({e}); trying scanned fallback")
            pairs = self._extract_pdf_scanned()

        return pairs

    def _process_pdf_page(self, page, page_num: int) -> List[ClinicalTablePair]:
        """Extract table-writeup pairs from one pdfplumber page."""
        pairs = []
        full_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
        tables = page.extract_tables() or []

        for t_idx, raw_table in enumerate(tables):
            if not raw_table or len(raw_table) < 2:
                continue

            linearized = self._linearize_raw_table(
                raw_table,
                title_hint=self._find_nearby_title(full_text, t_idx)
            )

            # Find writeup: text that follows the table on the same page
            writeup = self._find_writeup_in_text(full_text, linearized)

            if writeup:
                pairs.append(ClinicalTablePair(
                    table_text=linearized,
                    writeup=writeup,
                    source=self.source_path.name,
                    table_type=self._classify_table(linearized),
                    page=page_num,
                    pair_id=f"pdf_p{page_num}_t{t_idx+1}"
                ))

        return pairs

    def _extract_pdf_scanned(self) -> List[ClinicalTablePair]:
        """
        For scanned PDFs: render each page as an image using PyMuPDF,
        then delegate to ClinicalImageExtractor.
        """
        try:
            import fitz  # PyMuPDF
            from src.data_processing.image_extractor import ClinicalImageExtractor

            img_extractor = ClinicalImageExtractor(use_gpu=self.use_gpu_ocr)
            pairs = []
            doc = fitz.open(str(self.source_path))

            for page_num, page in enumerate(doc, 1):
                # Render at 200 DPI (sufficient for OCR)
                mat = fitz.Matrix(200 / 72, 200 / 72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_bytes = pix.tobytes("png")

                result = img_extractor.extract(img_bytes)
                for table in result.tables:
                    if table.n_rows >= 2:
                        pairs.append(ClinicalTablePair(
                            table_text=table.linearized,
                            writeup="",      # no paired writeup from scanned PDF
                            source=self.source_path.name,
                            table_type=self._classify_table(table.linearized),
                            page=page_num,
                            pair_id=f"scan_p{page_num}_{table.table_id}"
                        ))
            doc.close()
            return pairs

        except ImportError:
            logger.error("PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")
            return []

    # ─────────────────────────────────────────────────────────────────────────
    # DOCX Extraction
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_docx(self) -> List[ClinicalTablePair]:
        """
        Extract tables and adjacent paragraphs from a .docx file.
        python-docx gives direct access to table cells.
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return []

        doc = Document(str(self.source_path))
        pairs = []
        t_idx = 0

        # Iterate through document body elements in order
        # Elements are either paragraphs or tables
        body_elements = list(doc.element.body)
        all_blocks = doc.paragraphs + [None]  # sentinel

        # Use python-docx iterative block approach
        tables = doc.tables
        paragraphs = doc.paragraphs

        for t_idx, table in enumerate(tables):
            linearized = self._linearize_docx_table(table, t_idx)
            if not linearized:
                continue

            # Find the writeup: paragraph immediately after this table
            # (heuristic: use paragraph text that matches writeup pattern)
            writeup = self._find_writeup_in_docx(doc, t_idx)

            pairs.append(ClinicalTablePair(
                table_text=linearized,
                writeup=writeup,
                source=self.source_path.name,
                table_type=self._classify_table(linearized),
                page=0,
                pair_id=f"docx_t{t_idx+1}"
            ))

        return pairs

    def _linearize_docx_table(self, table, idx: int) -> str:
        """Convert python-docx Table to linearized string."""
        if not table.rows:
            return ""

        # Try to get title from first merged row (common in clinical tables)
        first_row_text = " | ".join(
            c.text.strip() for c in table.rows[0].cells
        )

        # Detect if first row is a title (no pipes pattern, single merged cell)
        unique_cells = list(set(c.text.strip() for c in table.rows[0].cells))
        if len(unique_cells) == 1 and unique_cells[0]:
            title = unique_cells[0]
            data_rows = table.rows[1:]
        else:
            title = f"Table {idx + 1}"
            data_rows = table.rows

        if len(data_rows) < 2:
            return ""

        # Build headers from first data row
        headers = [c.text.strip() for c in data_rows[0].cells]
        # Deduplicate merged header cells
        seen = set()
        clean_headers = []
        for h in headers:
            if h not in seen or not h:
                clean_headers.append(h)
                seen.add(h)

        parts = [f"start_table [TABLE_TITLE: {title}]"]
        parts.append(f"[HEADERS: | {' | '.join(clean_headers)}]")

        for row in data_rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            # Skip empty rows
            if all(not c for c in cells):
                continue
            # Deduplicate merged cells
            clean_cells = []
            prev = None
            for c in cells:
                if c != prev:
                    clean_cells.append(c)
                prev = c
            parts.append(f"[ROW] {' | '.join(clean_cells)}")

        parts.append("end_table")
        return " ".join(parts)

    def _find_writeup_in_docx(self, doc, table_idx: int) -> str:
        """
        Find the clinical narrative that follows a table in a docx.
        Looks at paragraphs after the table position in the document.
        """
        from docx.oxml.ns import qn

        body = doc.element.body
        children = list(body)

        # Find the table_idx-th w:tbl element
        tbl_count = 0
        tbl_pos = None
        for i, child in enumerate(children):
            if child.tag == qn('w:tbl'):
                if tbl_count == table_idx:
                    tbl_pos = i
                    break
                tbl_count += 1

        if tbl_pos is None:
            return ""

        # Collect paragraphs following the table
        writeup_parts = []
        for child in children[tbl_pos + 1:tbl_pos + 10]:
            if child.tag == qn('w:p'):
                text = "".join(n.text or "" for n in child.iter(qn('w:t')))
                if text.strip():
                    if self._WRITEUP_RE.search(text):
                        writeup_parts.append(text.strip())
                    elif writeup_parts:
                        writeup_parts.append(text.strip())
            elif child.tag == qn('w:tbl'):
                break  # hit next table

        return " ".join(writeup_parts)

    # ─────────────────────────────────────────────────────────────────────────
    # Image Extraction
    # ─────────────────────────────────────────────────────────────────────────

    def _extract_image(self) -> List[ClinicalTablePair]:
        """Delegate to ClinicalImageExtractor for image files."""
        from src.data_processing.image_extractor import ClinicalImageExtractor

        extractor = ClinicalImageExtractor(use_gpu=self.use_gpu_ocr)
        result = extractor.extract(self.source_path)

        pairs = []
        for table in result.tables:
            if table.n_rows >= 2:
                pairs.append(ClinicalTablePair(
                    table_text=table.linearized,
                    writeup="",  # no paired writeup from standalone image
                    source=self.source_path.name,
                    table_type=self._classify_table(table.linearized),
                    page=table.page,
                    pair_id=table.table_id
                ))
        return pairs

    # ─────────────────────────────────────────────────────────────────────────
    # Shared Helpers
    # ─────────────────────────────────────────────────────────────────────────

    def _linearize_raw_table(self, raw_table: List[List], title_hint: str = "") -> str:
        """Convert a pdfplumber raw table (list-of-lists) to linearized format."""
        if not raw_table:
            return ""

        title = title_hint or "Clinical Safety Table"
        parts = [f"start_table [TABLE_TITLE: {title}]"]

        headers = [str(c or "").strip() for c in raw_table[0]]
        parts.append(f"[HEADERS: | {' | '.join(headers)}]")

        for row in raw_table[1:]:
            if not any(c for c in row):
                continue
            cells = [str(c or "").strip() for c in row]
            parts.append(f"[ROW] {' | '.join(cells)}")

        parts.append("end_table")
        return " ".join(parts)

    def _find_nearby_title(self, full_text: str, table_idx: int) -> str:
        """Extract a table title from the surrounding page text."""
        pattern = re.compile(
            r'(?:Table|Figure)\s+\d+[\.:][^\n]{5,80}',
            re.IGNORECASE
        )
        matches = pattern.findall(full_text)
        if matches and table_idx < len(matches):
            return matches[table_idx].strip()
        return ""

    def _find_writeup_in_text(self, full_text: str, linearized: str) -> str:
        """
        Find the narrative writeup following the table in the full page text.
        Uses multi-pattern matching to find clinical summary paragraphs.
        """
        for pattern_str in self.WRITEUP_STARTERS:
            m = re.search(
                pattern_str + r'.{50,800}?(?=\n\n|\Z)',
                full_text,
                re.IGNORECASE | re.DOTALL
            )
            if m:
                candidate = m.group(0).strip()
                # Sanity: must contain at least one number to be a real writeup
                if re.search(r'\d', candidate):
                    return candidate

        return ""

    def _classify_table(self, linearized: str) -> str:
        """Classify table type for stratified splitting."""
        lin_lower = linearized.lower()
        if any(k in lin_lower for k in ("teae", "sae", "adverse event")):
            return "adverse_event"
        elif any(k in lin_lower for k in ("demographic", "age", "sex", "gender", "race")):
            return "demographics"
        elif any(k in lin_lower for k in ("efficacy", "response", "survival", "pfs", "os")):
            return "efficacy"
        elif any(k in lin_lower for k in ("laboratory", "lab", "haematology", "hematology")):
            return "laboratory"
        else:
            return "other"


# ─────────────────────────────────────────────────────────────────────────────
# Convenience functions for API use
# ─────────────────────────────────────────────────────────────────────────────

def extract_from_file(
    file_path: str,
    use_gpu_ocr: bool = False
) -> List[Dict]:
    """
    One-line extraction from any supported file (PDF, DOCX, JPEG, PNG).
    Returns list of dicts with 'table_text', 'writeup', 'table_type', etc.
    """
    extractor = ClinicalPDFExtractor(file_path, use_gpu_ocr=use_gpu_ocr)
    pairs = extractor.extract_all()
    return [p.to_dict() for p in pairs]


def extract_from_bytes(
    file_bytes: bytes,
    filename: str,
    use_gpu_ocr: bool = False
) -> List[Dict]:
    """
    Extract from raw file bytes (e.g., FastAPI UploadFile).
    Writes to a temp file then delegates to extract_from_file.
    """
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        return extract_from_file(tmp_path, use_gpu_ocr=use_gpu_ocr)
    finally:
        Path(tmp_path).unlink(missing_ok=True)


if __name__ == "__main__":
    import sys

    path = sys.argv[1] if len(sys.argv) > 1 else "data/raw/dataset.pdf"
    extractor = ClinicalPDFExtractor(path)
    pairs = extractor.extract_all()
    extractor.save("data/processed/")
    print(f"\n✅ Extracted {len(pairs)} pairs")
    if pairs:
        print(f"   First pair preview:")
        print(f"   Table: {pairs[0].table_text[:100]}...")
        if pairs[0].writeup:
            print(f"   Writeup: {pairs[0].writeup[:100]}...")
