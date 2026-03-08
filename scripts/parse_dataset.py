"""
Advanced Dataset Parser — Fixed for dada.docx structure.
Anchors on "Write up of Table X" markers.
"""
import json
import re
import argparse
from pathlib import Path
from typing import List, Dict

try:
    from docx import Document
except ImportError:
    pass


def parse_docx_structured(path: str) -> List[Dict]:
    doc = Document(path)
    pairs = []
    
    # 1. Linearize all Table objects in the doc
    linearized_tables = []
    for table in doc.tables:
        linearized_tables.append(linearize_docx_table(table))
    
    print(f"Found {len(linearized_tables)} Table objects in DOCX.")
    
    # 2. Extract all narrative summaries
    # We look for "Write up of Table X" followed by the next paragraph skiping empty ones.
    summaries = []
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if re.search(r'^Write\s*up\s*of\s*Table\s*\d+', text, re.IGNORECASE):
            # Collect ALL subsequent paragraphs until we hit the next Table or Write-up
            current_summary = []
            for j in range(i + 1, len(paragraphs)):
                p_text = paragraphs[j].text.strip()
                if not p_text:
                    continue
                # If we hit a new Table TITLE or another Write-up marker AT START, stop
                if re.search(r'^Table\s*\d+', p_text, re.IGNORECASE) or \
                   re.search(r'^Write\s*up\s*of\s*Table', p_text, re.IGNORECASE):
                    break
                current_summary.append(p_text)
            
            if current_summary:
                summaries.append(" ".join(current_summary))
            else:
                summaries.append("")
                    
    print(f"Found {len(summaries)} 'Write up' narrative paragraphs.")
    
    # 3. Pair them by order
    n = min(len(linearized_tables), len(summaries))
    for i in range(n):
        pairs.append({
            'table_text': linearized_tables[i],
            'writeup':    summaries[i]
        })
        
    return pairs


def linearize_docx_table(table) -> str:
    if not table.rows:
        return ''
    
    # Capture ALL text from cells to ensure no data is lost
    rows_text = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows_text.append(' | '.join(cells))
        
    if not rows_text:
        return ''
        
    header = rows_text[0]
    body = ' [ROW] '.join(rows_text[1:])
    
    return f'start_table [HEADERS: {header}] [ROW] {body} end_table'


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()

    pairs = parse_docx_structured(args.input)
    print(f"✅ Successfully paired {len(pairs)} table-writeup examples.")
    
    with open(args.output, 'w') as f:
        json.dump(pairs, f, indent=2)

if __name__ == '__main__':
    main()
