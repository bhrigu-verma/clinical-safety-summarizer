from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "data/eval_results/final_4mode_comparison_tier1_20260319_114452.csv"
OUT_PATH = ROOT / "PROFESSOR_PROGRESS_REPORT_MAR_2026.docx"

FIGURES = [
    (
        ROOT / "data/figures/figure1_architecture.png",
        "Figure 1. Hybrid ML+DL architecture with verification gate and fallback control.",
    ),
    (
        ROOT / "data/figures/figure16_4mode_faithfulness_20260319_120516.png",
        "Figure 2. Four-model faithfulness comparison (NAR, HR, OR, ACR views).",
    ),
    (
        ROOT / "data/figures/figure17_4mode_safety_fluency_20260319_120516.png",
        "Figure 3. Safety versus fluency behavior across all evaluated approaches.",
    ),
    (
        ROOT / "data/figures/figure18_4mode_frontier_20260319_120516.png",
        "Figure 4. Frontier view for quality-safety tradeoff and model selection.",
    ),
    (
        ROOT / "data/figures/figure19_4mode_latency_20260319_120516.png",
        "Figure 5. Latency profile (mean and upper-tail latency) across four approaches.",
    ),
]


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph("Clinical Safety Summarization")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(20)

    subtitle = doc.add_paragraph("Progress Report")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph("")
    author = doc.add_paragraph("Prepared by: Bhrigu Verma")
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    report_date = doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    report_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Missing figure: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    c = doc.add_paragraph(caption)
    c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    c.runs[0].italic = True


def load_mode_metrics() -> dict[str, dict[str, float]]:
    data: dict[str, dict[str, float]] = {}
    with CSV_PATH.open("r", newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            mode = row["mode"]
            data[mode] = {
                "rouge_l_mean": float(row["rouge_l_mean"]),
                "nar_mean": float(row["nar_mean"]),
                "hr_mean": float(row["hr_mean"]),
                "or_mean": float(row["or_mean"]),
                "acr_mean": float(row["acr_mean"]),
                "safety_score_mean": float(row["safety_score_mean"]),
                "fluency_score_mean": float(row["fluency_score_mean"]),
                "latency_ms_mean": float(row["latency_ms_mean"]),
                "latency_ms_p95": float(row["latency_ms_p95"]),
            }
    return data


def add_metrics_table(doc: Document, metrics: dict[str, dict[str, float]]) -> None:
    doc.add_paragraph("Table 1. Four-model comparison on Tier-1 benchmark (n=41)")
    table = doc.add_table(rows=1, cols=10)
    table.style = "Light List Accent 1"

    headers = [
        "Mode",
        "ROUGE-L",
        "NAR",
        "HR",
        "OR",
        "ACR",
        "Safety",
        "Fluency",
        "Latency Mean (ms)",
        "Latency P95 (ms)",
    ]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h

    mode_order = ["ml", "dl_base", "finetuned", "t5xxl"]
    mode_names = {
        "ml": "ML deterministic",
        "dl_base": "DL base",
        "finetuned": "DL fine-tuned",
        "t5xxl": "T5-XXL",
    }

    for mode in mode_order:
        if mode not in metrics:
            continue
        m = metrics[mode]
        row = table.add_row().cells
        row[0].text = mode_names[mode]
        row[1].text = f"{m['rouge_l_mean']:.4f}"
        row[2].text = f"{m['nar_mean']:.4f}"
        row[3].text = f"{m['hr_mean']:.4f}"
        row[4].text = f"{m['or_mean']:.4f}"
        row[5].text = f"{m['acr_mean']:.4f}"
        row[6].text = f"{m['safety_score_mean']:.4f}"
        row[7].text = f"{m['fluency_score_mean']:.4f}"
        row[8].text = f"{m['latency_ms_mean']:.2f}"
        row[9].text = f"{m['latency_ms_p95']:.2f}"


def build_doc() -> None:
    metrics = load_mode_metrics()

    doc = Document()
    set_doc_style(doc)
    add_title_page(doc)

    add_heading(doc, "1. Problem and Motivation", level=1)
    add_paragraphs(
        doc,
        [
            "Clinical safety reporting teams spend significant effort transforming structured adverse-event tables into regulatory narrative text. This conversion is repetitive, time-intensive, and vulnerable to inconsistency when performed manually at scale.",
            "A straightforward LLM solution is not sufficient for this setting. In regulated documentation, numeric hallucinations and arm-attribution swaps are unacceptable, even if the generated prose sounds fluent.",
            "The core project goal is therefore to automate table-to-narrative generation while preserving factual safety and maintaining readable, professional writing quality.",
        ],
    )

    add_heading(doc, "2. Solution Implemented", level=1)
    add_paragraphs(
        doc,
        [
            "We implemented a hybrid system that combines deterministic ML generation with neural rewriting. The deterministic layer ensures factual control, while the neural layer improves fluency.",
            "A verification gate checks generated output against source values and prevents unsafe generations from being used. If a violation is detected, the system falls back to the deterministic output.",
            "This architecture allows us to combine reliability and readability rather than trading one for the other.",
        ],
    )
    add_figure(doc, *FIGURES[0])

    add_heading(doc, "3. Approaches Tried", level=1)
    add_paragraphs(
        doc,
        [
            "To evaluate design choices rigorously, we compared four modes under the same Tier-1 benchmark conditions: (1) ML deterministic baseline, (2) DL base generation, (3) DL fine-tuned rewrite mode, and (4) T5-XXL comparison baseline.",
            "This setup gives clear evidence for what worked, what underperformed, and where the practical deployment tradeoffs appear.",
        ],
    )

    add_heading(doc, "4. Evaluation Setup", level=1)
    add_paragraphs(
        doc,
        [
            "All headline metrics in this report come from the consolidated four-model comparison file generated on 19 March 2026.",
            "We tracked faithfulness, safety, fluency, and latency to avoid making decisions from a single metric. This is important because high lexical overlap can mask safety-critical errors.",
        ],
    )
    add_metrics_table(doc, metrics)

    add_heading(doc, "5. Results and Interpretation", level=1)
    add_paragraphs(
        doc,
        [
            "The fine-tuned mode achieved the strongest ROUGE-L score while maintaining strong safety and fluency. The deterministic ML baseline remained the highest-safety profile and is still valuable as a fallback anchor.",
            "The base DL model showed weaker safety-quality balance than the fine-tuned configuration. T5-XXL, while large, was not practical in this setup due to poor benchmark behavior and significantly higher latency.",
            "Overall, the evidence supports the hybrid strategy as the most reliable path for regulated clinical summarization.",
        ],
    )

    for fig in FIGURES[1:]:
        add_figure(doc, *fig)

    add_heading(doc, "6. Work Completed to Date", level=1)
    completed = [
        "Built end-to-end data extraction, processing, and evaluation pipeline.",
        "Implemented deterministic ML generation and integrated it into the serving path.",
        "Integrated DL base and DL fine-tuned modes with safety verification.",
        "Established reproducible 4-model comparison workflow and generated publication-ready figures.",
        "Delivered frontend pathways for mode selection and summarization workflows.",
        "Produced engineering and executive documentation for architecture, experiments, and milestones.",
    ]
    for item in completed:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "7. Lessons Learned and Next Steps", level=1)
    add_paragraphs(
        doc,
        [
            "Key lesson: in regulated writing tasks, deterministic controls and verifiable checks are mandatory companions to neural generation.",
            "Next, we will improve fine-tuned mode safety margins, extend stress-test coverage, and finalize a paper-grade figure/table package for external presentation.",
            "We will also continue optimizing deployment practicality with a focus on latency stability and reproducible reporting outputs.",
        ],
    )

    doc.save(str(OUT_PATH))
    print(f"Generated: {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
